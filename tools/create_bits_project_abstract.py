from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path.home() / "Downloads" / "BITS_Project_Abstract_AI_Web_Test_Automation.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="B7C3D0", size="6"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(31, 31, 31)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(18)
    title.font.bold = True
    title.font.color.rgb = RGBColor(11, 37, 69)
    title.paragraph_format.space_after = Pt(8)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(12)
    subtitle.font.italic = True
    subtitle.font.color.rgb = RGBColor(70, 82, 95)
    subtitle.paragraph_format.space_after = Pt(14)

    for name, size, color in [
        ("Heading 1", 14, RGBColor(31, 77, 121)),
        ("Heading 2", 12, RGBColor(31, 77, 121)),
        ("Heading 3", 11, RGBColor(31, 77, 121)),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(5)


def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = "BITS Project Abstract"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = RGBColor(90, 100, 112)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("AI-Powered Intelligent Web Test Automation and Analytics Platform")
    footer_run.font.size = Pt(8.5)
    footer_run.font.color.rgb = RGBColor(90, 100, 112)


def add_title_block(doc):
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("AI-Powered Intelligent Web Test Automation and Analytics Platform")

    p = doc.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(
        "An AI-Driven System for Dynamic Test Generation, Automated Execution, "
        "and Failure Analysis of Web Applications"
    )

    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Inches(2.0), Inches(4.25)]
    fields = [
        ("Student Name", ""),
        ("Student ID / Roll No.", ""),
        ("Programme / Course", ""),
        ("Project Guide", ""),
        ("Semester", ""),
    ]
    for row, (label, value) in zip(table.rows, fields):
        row.cells[0].width = widths[0]
        row.cells[1].width = widths[1]
        set_cell_text(row.cells[0], label, bold=True)
        set_cell_text(row.cells[1], value)
        set_cell_shading(row.cells[0], "E8EEF5")
        for cell in row.cells:
            set_cell_border(cell)


def add_para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_scope_table(doc):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    headers = ["In Scope", "Out of Scope"]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, header, bold=True)
        set_cell_shading(cell, "D9EAF7")
        set_cell_border(cell)

    in_scope = [
        "FastAPI backend for test execution workflows",
        "Web interface for URL and prompt-based test execution",
        "Playwright integration for browser automation",
        "AI-assisted test scenario generation",
        "Dynamic website element extraction",
        "Execution logs, screenshots, and metrics capture",
        "PostgreSQL storage for test history",
        "Dashboards for pass/fail trends and analytics",
        "Downloadable test execution reports",
    ]
    out_scope = [
        "Mobile application testing",
        "Cross-browser cloud execution",
        "Distributed execution infrastructure",
        "Enterprise authentication and user management",
        "Full replacement of professional QA judgment",
    ]
    row = table.add_row()
    for idx, items in enumerate([in_scope, out_scope]):
        cell = row.cells[idx]
        cell.text = ""
        for item in items:
            p = cell.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            p.add_run(item)
        set_cell_border(cell)


def add_semester_plan(doc):
    doc.add_heading("Semester-Wise Work Plan", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    headers = ["Phase", "Duration", "Major Activities", "Deliverables"]
    widths = [Inches(0.95), Inches(1.05), Inches(3.0), Inches(1.55)]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = widths[i]
        set_cell_text(cell, header, bold=True)
        set_cell_shading(cell, "D9EAF7")
        set_cell_border(cell)

    rows = [
        (
            "Phase 1",
            "Weeks 1-3",
            "Requirement analysis, literature review, study of Playwright, FastAPI, AI-assisted test generation, and database requirements.",
            "Finalized scope, architecture outline, and project abstract.",
        ),
        (
            "Phase 2",
            "Weeks 4-6",
            "Backend setup, database schema design, API planning, and initial web interface design.",
            "FastAPI skeleton, PostgreSQL schema, and basic UI flow.",
        ),
        (
            "Phase 3",
            "Weeks 7-9",
            "Playwright automation integration, URL-based browser execution, screenshot capture, log collection, and basic test run storage.",
            "Working automation runner with execution history.",
        ),
        (
            "Phase 4",
            "Weeks 10-12",
            "AI prompt processing, test scenario generation, website element extraction, and structured test workflow creation.",
            "AI-assisted test generation module.",
        ),
        (
            "Phase 5",
            "Weeks 13-15",
            "Failure classification, AI-generated failure explanations, analytics dashboard, charts, and report generation.",
            "Dashboard, failure analysis, and downloadable reports.",
        ),
        (
            "Phase 6",
            "Weeks 16-18",
            "Testing, validation, documentation, final improvements, and project demonstration preparation.",
            "Final project report, demo, and presentation.",
        ),
    ]

    for row_data in rows:
        row = table.add_row()
        for i, text in enumerate(row_data):
            cell = row.cells[i]
            cell.width = widths[i]
            set_cell_text(cell, text)
            set_cell_border(cell)


def add_signatures(doc):
    doc.add_heading("Approval and Signatures", level=1)
    add_para(
        doc,
        "The following section is intentionally left blank for approval signatures "
        "and institutional verification."
    )
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    labels = [
        ("Student Signature", "Date"),
        ("Project Guide Signature", "Date"),
        ("Internal Examiner Signature", "Date"),
        ("Head / Coordinator Signature", "Date"),
    ]
    for row, pair in zip(table.rows, labels):
        for idx, label in enumerate(pair):
            cell = row.cells[idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(18)
            p.add_run(label + "\n\n")
            line = p.add_run("____________________________")
            line.font.color.rgb = RGBColor(120, 120, 120)
            set_cell_border(cell, color="DADCE0", size="4")


def build():
    doc = Document()
    style_document(doc)
    add_header_footer(doc)
    add_title_block(doc)

    doc.add_heading("1. Broad Area of Work", level=1)
    add_para(
        doc,
        "This project lies at the intersection of software testing, web automation, "
        "artificial intelligence, and full-stack application development. The work "
        "focuses on building an intelligent testing platform where a user can enter "
        "a website URL and natural language testing instructions. The system will "
        "generate relevant test scenarios, execute them through browser automation, "
        "record execution evidence, identify failures, and present actionable insights "
        "through analytics dashboards."
    )
    add_bullets(
        doc,
        [
            "Web application automation using Playwright",
            "Artificial intelligence for test case generation and failure analysis",
            "Backend development using FastAPI",
            "Database design using PostgreSQL and SQLAlchemy",
            "Reporting and visualization through interactive dashboards",
            "Human-computer interaction through natural language driven testing",
        ],
    )

    doc.add_heading("2. Background", level=1)
    add_para(
        doc,
        "Modern web applications change frequently and require continuous testing "
        "to maintain reliability, usability, and functional correctness. Conventional "
        "automation frameworks are powerful, but they usually require testers or "
        "developers to manually write scripts, update selectors, maintain test cases, "
        "and interpret logs after each execution."
    )
    add_para(
        doc,
        "Tools such as Playwright Codegen reduce some scripting effort by recording "
        "user actions, but they still depend on manual intervention and technical "
        "expertise. This creates difficulty for non-technical users and increases "
        "maintenance overhead when websites evolve quickly. Recent progress in Large "
        "Language Models creates an opportunity to assist this process by translating "
        "natural language intent into structured testing workflows and summarizing "
        "failures in a readable form."
    )

    doc.add_heading("3. Problem Statement", level=1)
    add_para(
        doc,
        "The problem addressed by this project is the high manual effort required "
        "to create, execute, and analyze web application test cases. Existing tools "
        "often separate test planning, automation execution, logging, and reporting "
        "into different workflows. As a result, teams spend significant time writing "
        "scripts and interpreting results instead of focusing on application quality."
    )
    add_para(
        doc,
        "The proposed system aims to reduce this effort by creating an integrated "
        "platform that uses AI to assist dynamic test generation and execution, while "
        "Playwright performs browser-level automation and the application records "
        "evidence for later review."
    )

    doc.add_heading("4. Proposed System", level=1)
    add_para(
        doc,
        "The proposed platform will accept a target website URL and natural language "
        "instructions from the user. The AI module will convert the instruction into "
        "structured test scenarios and steps. Playwright will execute these steps in "
        "a browser environment, while the backend captures logs, screenshots, execution "
        "status, and timing information. The system will classify outcomes as pass, "
        "fail, or warning, and provide AI-generated failure explanations where possible."
    )
    add_numbered(
        doc,
        [
            "User submits a website URL and testing prompt.",
            "The system extracts relevant page context and elements.",
            "AI generates a structured test workflow.",
            "Playwright executes the workflow in an automated browser session.",
            "Execution logs, screenshots, and metrics are stored.",
            "AI assists in summarizing failures and likely causes.",
            "Dashboards display execution trends, statistics, and reports.",
        ],
    )

    doc.add_heading("5. Objectives", level=1)
    add_bullets(
        doc,
        [
            "To develop a web-based testing platform capable of accepting website URLs and testing instructions.",
            "To integrate Playwright for browser automation and dynamic web interaction.",
            "To generate test scenarios from natural language prompts using AI models.",
            "To execute generated test workflows automatically.",
            "To capture execution logs, screenshots, and performance metrics.",
            "To classify test outcomes into pass, fail, and warning categories.",
            "To provide AI-generated explanations for failures.",
            "To visualize execution results through dashboards and charts.",
            "To store execution history for future analysis and reporting.",
        ],
    )

    doc.add_heading("6. Scope of Work", level=1)
    add_para(
        doc,
        "The scope is limited to AI-assisted web application testing through browser "
        "automation. The system is intended to assist test generation and analysis; "
        "it does not claim to test any website completely or replace expert validation."
    )
    add_scope_table(doc)

    doc.add_heading("7. Methodology", level=1)
    add_para(
        doc,
        "The project will follow an iterative engineering approach. First, the core "
        "backend, database, and browser automation runner will be established. Next, "
        "AI-assisted test generation will be integrated with structured prompts and "
        "page context extraction. The final stages will focus on failure analysis, "
        "dashboard development, report generation, testing, and documentation."
    )

    doc.add_heading("8. Technology Stack", level=1)
    add_bullets(
        doc,
        [
            "Frontend: React or equivalent web interface for test submission and dashboards.",
            "Backend: FastAPI for APIs, orchestration, and execution management.",
            "Automation: Playwright for browser control, screenshots, and interaction.",
            "Database: PostgreSQL with SQLAlchemy for structured execution history.",
            "AI Layer: Large Language Model based prompt interpretation, test generation, and failure summarization.",
            "Reporting: Dashboard charts and downloadable reports for test outcomes.",
        ],
    )

    add_semester_plan(doc)

    doc.add_heading("10. Expected Outcomes", level=1)
    add_bullets(
        doc,
        [
            "A working web platform for URL and prompt-based test execution.",
            "Automated Playwright execution with screenshots and logs.",
            "AI-assisted test scenario generation from natural language prompts.",
            "Failure summaries that explain likely causes in user-friendly language.",
            "Persistent test execution history stored in PostgreSQL.",
            "Analytics dashboards showing pass/fail statistics and trends.",
            "Downloadable reports suitable for review and documentation.",
        ],
    )

    doc.add_heading("11. Academic and Engineering Significance", level=1)
    add_para(
        doc,
        "The project is stronger than a conventional automation script because it "
        "combines AI, browser automation, backend engineering, database design, and "
        "analytics into a single system. It demonstrates practical software engineering "
        "while also exploring a current research direction: how AI can assist testing "
        "workflows without overstating its reliability or autonomy."
    )

    doc.add_heading("12. Conclusion", level=1)
    add_para(
        doc,
        "This project proposes an AI-powered intelligent web test automation and "
        "analytics platform that assists users in generating, executing, analyzing, "
        "and reporting web application tests. By combining natural language driven "
        "test creation with Playwright automation and analytics dashboards, the system "
        "aims to reduce manual effort, improve visibility into failures, and provide "
        "a realistic semester-scale engineering solution."
    )

    add_signatures(doc)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build()
