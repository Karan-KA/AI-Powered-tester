from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    Image as RLImage,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)


DOWNLOADS = Path.home() / "Downloads"
DOCX_OUT = DOWNLOADS / "abstract_karan2026_submission_ready_final.docx"
PDF_OUT = DOWNLOADS / "abstract_karan2026_submission_ready_final.pdf"
IMAGE_DIR = DOWNLOADS / "karan_abstract_images"
LOGO = IMAGE_DIR / "page1_1_Image28.png"
CONTENTS_IMAGE = IMAGE_DIR / "page2_1_Image38.png"
STUDENT_SIGN = IMAGE_DIR / "page9_1_Image63.jpg"
SUPERVISOR_SIGN = IMAGE_DIR / "page9_2_Image64.jpg"
EXAMINER_SIGN = IMAGE_DIR / "page9_3_Image65.jpg"


TITLE = "AI-Powered Intelligent Web Test Automation and Analytics Platform"
SUBTITLE = (
    "An AI-Driven System for Dynamic Test Generation, Automated Execution, "
    "and Failure Analysis of Web Applications"
)


metadata = [
    ("Course Title", "Dissertation / Project work"),
    ("Student Name", "Karan K A"),
    ("BITS ID", "202217b2235"),
    ("Program", "BSc Designs and Computing"),
    (
        "Research Area",
        "AI-driven software testing and intelligent web automation for dynamic test "
        "generation, automated execution, and failure analysis of web applications.",
    ),
    ("Project Carried Out at", "HCL Tech, Bengaluru"),
]


sections = [
    (
        "1. Broad Area of Work",
        [
            (
                "p",
                "This project is based on the combined fields of software quality assurance, "
                "browser-based automation, artificial intelligence, and web application "
                "development. The main focus is to design and develop an intelligent web "
                "testing platform that allows users to provide a website link along with "
                "testing requirements in natural language. Based on the given input, the "
                "system will assist in creating test scenarios, run them automatically in "
                "a browser environment, collect execution details, detect possible errors, "
                "and display useful results through an analytics dashboard.",
            ),
            ("p", "The project covers the following technical areas:"),
            (
                "bullets",
                [
                    "Browser automation using Playwright",
                    "AI-assisted generation of test cases and failure summaries",
                    "Backend API development using FastAPI",
                    "Data storage and management using PostgreSQL and SQLAlchemy",
                    "Result analysis and visualization through dashboards",
                    "Natural language based interaction for web application testing",
                ],
            ),
        ],
    ),
    (
        "2. Background",
        [
            (
                "p",
                "Web applications are updated regularly to add new features, improve user "
                "experience, and fix issues. Because of these frequent changes, continuous "
                "testing is necessary to ensure that the application works correctly, remains "
                "user-friendly, and performs reliably. Traditional automation tools can help "
                "in testing, but they often need skilled testers or developers to write scripts, "
                "modify element locators, maintain test cases, and study execution logs manually.",
            ),
            (
                "p",
                "Although tools like Playwright Codegen can record user actions and convert "
                "them into automation steps, they still require human review, technical "
                "knowledge, and regular maintenance. This makes automated testing difficult "
                "for users with limited programming experience. With the advancement of Large "
                "Language Models, there is now a possibility to make testing more accessible "
                "by converting natural language instructions into structured test workflows "
                "and generating simple explanations for test failures.",
            ),
        ],
    ),
    (
        "3. Problem Statement",
        [
            (
                "p",
                "Creating and maintaining automated tests for web applications usually "
                "requires considerable manual work. Testers must design test cases, write "
                "automation scripts, run them, check logs, capture evidence, and prepare "
                "reports. In many existing systems, these activities are handled through "
                "separate tools, which makes the overall testing process time-consuming "
                "and difficult to manage.",
            ),
            (
                "p",
                "This project addresses that gap by proposing a single platform that can "
                "assist in test creation, execution, result tracking, and failure understanding. "
                "The system will use AI to support dynamic test workflow generation from user "
                "instructions, while Playwright will handle browser automation. The platform "
                "will also save execution evidence such as logs, screenshots, and test outcomes "
                "for future review.",
            ),
        ],
    ),
    (
        "4. Proposed System",
        [
            (
                "p",
                "The proposed system is an AI-assisted web testing platform where the user "
                "provides a website URL and describes the required test in simple language. "
                "The AI component will interpret the user's instruction and prepare a "
                "structured set of test steps. These steps will then be executed using "
                "Playwright in an automated browser session. During execution, the backend "
                "will record important details such as logs, screenshots, status, errors, "
                "and execution time.",
            ),
            (
                "p",
                "The system will organize the result as passed, failed, or warning-based "
                "outcomes. When a failure occurs, the AI module will help generate a readable "
                "explanation so that the user can understand the possible reason without "
                "manually studying complex logs.",
            ),
            ("p", "Process flow:"),
            (
                "numbered",
                [
                    "The user enters the website URL and testing requirement.",
                    "The application reads the page structure and important elements.",
                    "The AI module prepares a structured testing workflow.",
                    "Playwright runs the generated steps in a browser.",
                    "Logs, screenshots, metrics, and results are saved.",
                    "The AI module summarizes failures and possible causes.",
                    "Dashboards show reports, trends, and execution statistics.",
                ],
            ),
        ],
    ),
    (
        "5. Objectives",
        [
            (
                "bullets",
                [
                    "To design and develop a web-based platform for AI-assisted web application testing.",
                    "To allow users to submit website URLs and testing requirements through a simple interface.",
                    "To use Playwright for automated browser interaction and test execution.",
                    "To generate structured test scenarios from natural language instructions.",
                    "To run the generated test workflows automatically.",
                    "To collect screenshots, logs, execution status, and performance-related details.",
                    "To categorize test results into pass, fail, and warning states.",
                    "To provide simple AI-generated explanations for failed test cases.",
                    "To display test results and trends using charts and dashboards.",
                    "To maintain previous test execution records for review, comparison, and reporting.",
                ],
            )
        ],
    ),
]


plan_rows = [
    (
        "Literature Review & Project Outline",
        "Week 1 - Week 2",
        "Study existing web automation tools, Playwright, AI-assisted testing methods, "
        "FastAPI, PostgreSQL, and related research work; prepare and submit the project outline.",
    ),
    (
        "System Design & Architecture",
        "Week 3 - Week 4",
        "Design the system architecture, database schema, API structure, user workflow, "
        "and interaction between frontend, backend, AI module, and Playwright automation.",
    ),
    (
        "Backend & Database Module",
        "Week 5 - Week 6",
        "Implement the FastAPI backend, create PostgreSQL tables, configure SQLAlchemy "
        "models, and build APIs for storing test execution records.",
    ),
    (
        "Playwright Automation Module",
        "Week 7 - Week 8",
        "Integrate Playwright for browser automation; implement URL-based browser launch, "
        "page interaction, screenshot capture, execution logs, and test status tracking.",
    ),
    (
        "AI Test Generation Module",
        "Week 9 - Week 10",
        "Develop the AI module to convert natural language testing instructions into "
        "structured test scenarios and executable workflow steps.",
    ),
    (
        "Failure Analysis & Reporting",
        "Week 11 - Week 12",
        "Implement pass, fail, and warning classification; generate AI-assisted failure "
        "explanations; summarize logs; map screenshots; and create downloadable reports.",
    ),
    (
        "Dashboard & Analytics UI",
        "Week 13 - Week 14",
        "Build the frontend dashboard to display test history, pass/fail statistics, "
        "execution trends, screenshots, logs, and report views.",
    ),
    (
        "Testing, Bug Fixing & Submission",
        "Week 15 - Week 16",
        "Perform end-to-end testing, fix edge cases, improve UI and backend reliability, "
        "finalize documentation, prepare the project report, and submit for review.",
    ),
]


references = [
    "Beizer, B. (1990). Software Testing Techniques (2nd ed.). Van Nostrand Reinhold. -- Foundational reference for understanding software testing principles, test case design, defect detection, and the importance of structured testing.",
    "W3C. (2018). WebDriver is now a W3C Recommendation. https://www.w3.org/news/2018/webdriver-is-now-a-w3c-recommendation/ -- Relevant for understanding browser automation standards and how automated tools control browsers for web application testing.",
    "Microsoft Playwright. (n.d.). Playwright Documentation. https://playwright.dev/docs/intro -- Official reference for Playwright-based browser automation, end-to-end testing, screenshots, tracing, and web interaction workflows used in this project.",
    "Vaswani, A., Shazeer, N., Parmar, N., Uszkoreit, J., Jones, L., Gomez, A. N., Kaiser, L., & Polosukhin, I. (2017). Attention Is All You Need. https://arxiv.org/abs/1706.03762 -- Foundational paper behind Transformer-based language models, which supports the AI component used for natural language interpretation and test workflow generation.",
    "Celik, A., & Mahmoud, Q. H. (2025). A Review of Large Language Models for Automated Test Case Generation. Machine Learning and Knowledge Extraction, 7(3), 97. https://www.mdpi.com/2504-4990/7/3/97 -- Provides research context on how Large Language Models are being used for automated test case generation.",
    "FastAPI. (n.d.). FastAPI Documentation. https://fastapi.tiangolo.com/ -- Official reference for developing high-performance backend APIs, request handling, and service endpoints used in the proposed testing platform.",
    "SQLAlchemy. (n.d.). SQLAlchemy Documentation. https://docs.sqlalchemy.org/ -- Reference for ORM-based database interaction, schema modeling, and managing test execution records in PostgreSQL.",
    "PostgreSQL Global Development Group. (n.d.). PostgreSQL Documentation. https://www.postgresql.org/docs/ -- Official documentation for relational database storage, queries, and data management used for maintaining test run history and reports.",
]


supervisor_table = [
    ["", "Supervisor", "Additional Examiner"],
    ["Name", "Gowdeswaran P C", "Aakash R"],
    ["Qualification", "BE ECE", "BE ECE"],
    ["Designation", "Senior technical lead", "Lead engineer"],
    ["Organisation and location", "HCL Tech, Bengaluru", "HCL Tech, Bengaluru"],
    ["Phone number", "9535099440", "9965647613"],
    ["Email Address", "gowdeswaran.pc@hcltech.com", "aakash-r@hcltech.com"],
]


remarks = [
    "This project presents a relevant and timely problem in the field of web application testing. As modern websites are updated regularly, maintaining test cases and automation scripts becomes a difficult task for developers and testers. The proposed idea of accepting a website URL and natural language testing instructions makes the system practical, user-friendly, and suitable for real-world testing needs.",
    "The planned solution is technically sound and achievable within the project duration. Playwright is an appropriate choice for browser automation, while the AI-based module adds an intelligent layer for generating test workflows and explaining failures. The inclusion of FastAPI, PostgreSQL, execution logs, screenshots, dashboards, and reports shows that the project covers both implementation and analysis aspects.",
    "The project has good academic value because it focuses on AI-assisted testing rather than simple script-based automation. It attempts to reduce manual effort in test creation, execution, and result interpretation. This makes the work suitable for a software engineering research project with practical application.",
    "Overall, the project is well scoped and has strong potential. The student should concentrate on developing a dependable prototype that demonstrates prompt-based test generation, automated browser execution, failure classification, evidence collection, and analytics-based reporting.",
]


def set_table_cell_border(cell, color="C9C9C9", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, size=11, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_docx(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(7)

    title = doc.styles["Title"]
    title.font.name = "Times New Roman"
    title.font.size = Pt(18)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0, 0, 0)

    subtitle = doc.styles["Subtitle"]
    subtitle.font.name = "Times New Roman"
    subtitle.font.size = Pt(13)
    subtitle.font.bold = False
    subtitle.font.italic = False
    subtitle.font.color.rgb = RGBColor(0, 0, 0)

    for style_name, size in [("Heading 1", 16), ("Heading 2", 14)]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(14)
        style.paragraph_format.space_after = Pt(8)


def add_docx_para(doc, text, justify=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(text)


def add_docx_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_docx_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def build_docx():
    doc = Document()
    style_docx(doc)

    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO), width=Inches(0.85))

    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(TITLE)
    p = doc.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(SUBTITLE)
    doc.add_paragraph()

    table = doc.add_table(rows=len(metadata), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for idx, (label, value) in enumerate(metadata):
        row = table.rows[idx]
        row.cells[0].width = Inches(2.2)
        row.cells[1].width = Inches(4.7)
        set_cell_text(row.cells[0], label, bold=True)
        set_cell_text(row.cells[1], value)
        for cell in row.cells:
            set_table_cell_border(cell, color="FFFFFF", size="0")
    doc.add_paragraph()

    institute = doc.add_paragraph()
    institute.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = institute.add_run("BIRLA INSTITUTE OF TECHNOLOGY & SCIENCE, PILANI")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)

    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BIRLA INSTITUTE OF TECHNOLOGY & SCIENCE, PILANI")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    doc.add_heading("Contents", level=1)
    if CONTENTS_IMAGE.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(CONTENTS_IMAGE), width=Inches(5.8))
    doc.add_page_break()

    for idx, (heading, blocks) in enumerate(sections, 1):
        doc.add_heading(heading, level=1)
        for block_type, value in blocks:
            if block_type == "p":
                add_docx_para(doc, value)
            elif block_type == "bullets":
                add_docx_bullets(doc, value)
            elif block_type == "numbered":
                add_docx_numbered(doc, value)

    doc.add_heading("Plan of Work", level=1)
    plan_table = doc.add_table(rows=1, cols=3)
    plan_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Phase", "Start - End Date", "Work to be Done"]
    for i, header in enumerate(headers):
        cell = plan_table.rows[0].cells[i]
        set_cell_text(cell, header, bold=True, size=12, color=(255, 255, 255))
        shade_cell(cell, "4F4F4F")
        set_table_cell_border(cell, color="777777")
    for phase, dates, work in plan_rows:
        row = plan_table.add_row()
        for i, text in enumerate([phase, dates, work]):
            set_cell_text(row.cells[i], text, size=11)
            set_table_cell_border(row.cells[i])

    doc.add_heading("6. Literature References", level=1)
    add_docx_bullets(doc, references)

    doc.add_heading("7. Particulars of the Supervisor and Examiner", level=1)
    sup_table = doc.add_table(rows=len(supervisor_table), cols=3)
    sup_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row_data in enumerate(supervisor_table):
        for c_idx, text in enumerate(row_data):
            cell = sup_table.rows[r_idx].cells[c_idx]
            set_cell_text(cell, text, bold=(r_idx == 0 or c_idx == 0), size=10)
            if r_idx == 0:
                shade_cell(cell, "E6E6E6")
            set_table_cell_border(cell)

    doc.add_heading("8. Remarks of the Supervisor", level=1)
    for text in remarks:
        add_docx_para(doc, text)

    doc.add_paragraph()
    sig_table = doc.add_table(rows=2, cols=3)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sigs = ["Signature of the student", "Signature of the supervisor", "Signature of the Additional Examiner"]
    names = ["Name: Karan K A", "Name: Gowdeswaran P C", "Name: Aakash R"]
    sign_images = [STUDENT_SIGN, SUPERVISOR_SIGN, EXAMINER_SIGN]
    for c in range(3):
        cell = sig_table.rows[0].cells[c]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if sign_images[c].exists():
            p.add_run().add_picture(str(sign_images[c]), width=Inches(1.25))
        p.add_run("\n" + sigs[c])
        set_cell_text(sig_table.rows[1].cells[c], names[c], size=11)
        for r in range(2):
            set_table_cell_border(sig_table.rows[r].cells[c], color="FFFFFF", size="0")

    doc.save(DOCX_OUT)


def pdf_footer(canvas, doc):
    canvas.saveState()
    canvas.setFont("Times-Roman", 8)
    canvas.drawString(0.45 * inch, 0.35 * inch, "Classification: Internal")
    canvas.restoreState()


def get_pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="CoverTitle",
            parent=styles["Title"],
            fontName="Times-Bold",
            fontSize=18,
            leading=23,
            alignment=TA_CENTER,
            spaceAfter=10,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CoverSubtitle",
            parent=styles["Normal"],
            fontName="Times-Roman",
            fontSize=13,
            leading=17,
            alignment=TA_CENTER,
            spaceAfter=14,
        )
    )
    styles.add(
        ParagraphStyle(
            name="SectionHeading",
            parent=styles["Heading1"],
            fontName="Times-Bold",
            fontSize=16,
            leading=20,
            textColor=colors.black,
            spaceBefore=9,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="BodyJustify",
            parent=styles["BodyText"],
            fontName="Times-Roman",
            fontSize=11.5,
            leading=15,
            alignment=TA_LEFT,
            spaceAfter=5,
        )
    )
    styles.add(
        ParagraphStyle(
            name="BodyCenter",
            parent=styles["BodyText"],
            fontName="Times-Roman",
            fontSize=11.5,
            leading=15,
            alignment=TA_CENTER,
            spaceAfter=5,
        )
    )
    styles.add(
        ParagraphStyle(
            name="BulletText",
            parent=styles["BodyText"],
            fontName="Times-Roman",
            fontSize=11.5,
            leading=14,
            leftIndent=18,
            firstLineIndent=-10,
            spaceAfter=1,
        )
    )
    styles.add(
        ParagraphStyle(
            name="SmallTable",
            parent=styles["BodyText"],
            fontName="Times-Roman",
            fontSize=9.6,
            leading=12,
            alignment=TA_LEFT,
        )
    )
    styles.add(
        ParagraphStyle(
            name="SmallTableBold",
            parent=styles["SmallTable"],
            fontName="Times-Bold",
        )
    )
    return styles


def para(text, style):
    safe = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    return Paragraph(safe, style)


def bullet_para(text, style):
    safe = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    return Paragraph(f"-&nbsp;&nbsp;{safe}", style)


def build_pdf():
    styles = get_pdf_styles()
    doc = BaseDocTemplate(
        str(PDF_OUT),
        pagesize=A4,
        leftMargin=0.58 * inch,
        rightMargin=0.58 * inch,
        topMargin=0.62 * inch,
        bottomMargin=0.62 * inch,
    )
    frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id="normal")
    doc.addPageTemplates([PageTemplate(id="main", frames=[frame], onPage=pdf_footer)])
    story = []

    story.append(Spacer(1, 0.2 * inch))
    if LOGO.exists():
        logo = RLImage(str(LOGO), width=0.82 * inch, height=0.79 * inch)
        logo.hAlign = "CENTER"
        story.append(logo)
        story.append(Spacer(1, 0.16 * inch))
    story.append(Paragraph(TITLE, styles["CoverTitle"]))
    story.append(Paragraph(SUBTITLE, styles["CoverSubtitle"]))
    meta_data = [[para(label, styles["SmallTableBold"]), para(value, styles["SmallTable"])] for label, value in metadata]
    table = Table(meta_data, colWidths=[1.9 * inch, 4.55 * inch], hAlign="CENTER")
    table.setStyle(
        TableStyle(
            [
                ("FONT", (0, 0), (-1, -1), "Times-Roman", 11),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    story.append(table)
    story.append(Spacer(1, 0.4 * inch))
    story.append(Paragraph("<b>BIRLA INSTITUTE OF TECHNOLOGY &amp; SCIENCE, PILANI</b>", styles["BodyCenter"]))
    story.append(PageBreak())

    story.append(Paragraph("<b>BIRLA INSTITUTE OF TECHNOLOGY &amp; SCIENCE, PILANI</b>", styles["BodyCenter"]))
    story.append(Spacer(1, 0.25 * inch))
    story.append(Paragraph("Contents", styles["SectionHeading"]))
    if CONTENTS_IMAGE.exists():
        toc_img = RLImage(str(CONTENTS_IMAGE), width=6.6 * inch, height=3.55 * inch)
        toc_img.hAlign = "CENTER"
        story.append(Spacer(1, 0.15 * inch))
        story.append(toc_img)
    story.append(PageBreak())

    for idx, (heading, blocks) in enumerate(sections, 1):
        story.append(Paragraph(heading, styles["SectionHeading"]))
        for block_type, value in blocks:
            if block_type == "p":
                story.append(para(value, styles["BodyJustify"]))
            elif block_type == "bullets":
                for item in value:
                    story.append(bullet_para(item, styles["BulletText"]))
            elif block_type == "numbered":
                for i, item in enumerate(value, 1):
                    story.append(para(f"{i}. {item}", styles["BulletText"]))

    story.append(Paragraph("Plan of Work", styles["SectionHeading"]))
    plan_data = [[para("Phase", styles["SmallTableBold"]), para("Start - End Date", styles["SmallTableBold"]), para("Work to be Done", styles["SmallTableBold"])]]
    for row in plan_rows:
        plan_data.append([para(row[0], styles["SmallTable"]), para(row[1], styles["SmallTable"]), para(row[2], styles["SmallTable"])])
    plan = Table(plan_data, colWidths=[1.65 * inch, 1.35 * inch, 4.0 * inch], repeatRows=1, hAlign="LEFT")
    plan.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4F4F4F")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#BFBFBF")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    story.append(plan)

    story.append(Paragraph("6. Literature References", styles["SectionHeading"]))
    for ref in references:
        story.append(bullet_para(ref, styles["BulletText"]))

    story.append(Paragraph("7. Particulars of the Supervisor and Examiner", styles["SectionHeading"]))
    sup_data = [[para(cell, styles["SmallTableBold"] if r == 0 or c == 0 else styles["SmallTable"]) for c, cell in enumerate(row)] for r, row in enumerate(supervisor_table)]
    sup = Table(sup_data, colWidths=[1.65 * inch, 2.6 * inch, 2.7 * inch], repeatRows=1, hAlign="LEFT")
    sup.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E6E6E6")),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#BFBFBF")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    story.append(sup)

    story.append(Paragraph("8. Remarks of the Supervisor", styles["SectionHeading"]))
    for text in remarks:
        story.append(para(text, styles["BodyJustify"]))

    story.append(Spacer(1, 0.35 * inch))
    def signature_cell(path, label):
        if path.exists():
            return [RLImage(str(path), width=1.25 * inch, height=0.55 * inch), para(label, styles["SmallTable"])]
        return [para(label, styles["SmallTable"])]

    sig_data = [
        [
            signature_cell(STUDENT_SIGN, "Signature of the student"),
            signature_cell(SUPERVISOR_SIGN, "Signature of the supervisor"),
            signature_cell(EXAMINER_SIGN, "Signature of the Additional Examiner"),
        ],
        [para("Name: Karan K A", styles["SmallTable"]), para("Name: Gowdeswaran P C", styles["SmallTable"]), para("Name: Aakash R", styles["SmallTable"])],
    ]
    sig = Table(sig_data, colWidths=[2.1 * inch, 2.3 * inch, 2.5 * inch], hAlign="LEFT")
    sig.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "TOP"), ("TOPPADDING", (0, 0), (-1, -1), 12)]))
    story.append(sig)

    doc.build(story)


if __name__ == "__main__":
    build_docx()
    build_pdf()
