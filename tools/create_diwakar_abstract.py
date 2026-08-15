from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    Image as RLImage,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)


DOWNLOADS = Path.home() / "Downloads"
IMAGE_DIR = DOWNLOADS / "karan_abstract_images"
LOGO = IMAGE_DIR / "page1_1_Image28.png"
CONTENTS_IMAGE = IMAGE_DIR / "page2_1_Image38.png"
PDF_OUT = DOWNLOADS / "diwakar_ok.pdf"
DOCX_OUT = DOWNLOADS / "diwakar_ok.docx"
SIGN_DIR = DOWNLOADS / "diwakar_signatures"
SUPERVISOR_SIGN = SIGN_DIR / "supervisor_clean.jpg"
EXAMINER_SIGN = SIGN_DIR / "examiner_clean.jpg"
STUDENT_SIGN = SIGN_DIR / "student_clean.jpg"


TITLE = "Secure Offline Log Analysis using Local LLM (AI)"
SUBTITLE = "An AI-Driven System for Local Log Processing, Root Cause Analysis, and Debugging Recommendations"


metadata = [
    ("Course Title", "Dissertation / Project work"),
    ("Student Name", "Diwakar S"),
    ("BITS ID", "202217b2100"),
    ("Program", "BSc - Design and Computing-HCL"),
    (
        "Research Area",
        "Artificial Intelligence",
    ),
    ("Project Carried Out at", "HCLTech"),
]


sections = [
    (
        "1. Broad Area of Work",
        [
            (
                "p",
                "This project is based on the combined areas of log processing, artificial intelligence, local large language model deployment, and web application development. The main objective is to design and develop an AI-powered log analyzer that allows users to upload system or application logs and receive structured analysis, error summaries, possible root causes, and suggested corrective actions.",
            ),
            ("p", "The project covers the following technical areas:"),
            (
                "bullets",
                [
                    "Log parsing and pattern extraction using Python",
                    "Local LLM integration using Ollama",
                    "Backend API development using Flask",
                    "Frontend development using HTML, CSS, React, and JavaScript",
                    "Retrieval-augmented generation concepts for known issue matching",
                    "Visualization of errors, warnings, and log trends through dashboards",
                ],
            ),
        ],
    ),
    (
        "2. Background",
        [
            (
                "p",
                "Modern software systems, drivers, services, and validation environments generate large volumes of logs during execution. These logs contain useful information such as errors, warnings, timestamps, component failures, service restarts, and performance-related events. However, manually reading long log files and identifying the exact cause of a failure is time-consuming and requires technical experience.",
            ),
            (
                "p",
                "Traditional log analysis methods depend on keyword search, manual filtering, and predefined rules. These approaches are useful but limited when the log contains multiple related events or when the user needs a clear explanation of the failure. Recent progress in local large language models provides an opportunity to analyze logs privately on the user's own system and generate readable summaries without sending sensitive data to external cloud services.",
            ),
        ],
    ),
    (
        "3. Problem Statement",
        [
            (
                "p",
                "The problem addressed by this project is the difficulty involved in analyzing large and complex log files manually. Developers, testers, and validation engineers often spend significant time searching for error patterns, comparing events, understanding crash messages, and preparing debugging notes. This process becomes harder when logs are lengthy, noisy, or generated from multiple system components.",
            ),
            (
                "p",
                "The proposed system aims to reduce this effort by creating an integrated log analysis platform that uses a local LLM to summarize errors, classify issues, identify possible root causes, and suggest practical debugging steps while keeping the log data within the local environment.",
            ),
        ],
    ),
    (
        "4. Proposed System",
        [
            (
                "p",
                "The proposed system will allow the user to upload a log file through a web interface. The backend will parse the uploaded file, extract important events such as ERROR, WARNING, INFO, and CRITICAL messages, and prepare structured input for the local LLM. Ollama will be used to run the selected model locally and generate analysis based on the extracted log context.",
            ),
            (
                "p",
                "The platform will display detected issues, likely causes, suggested actions, and summarized findings through a clean dashboard. In an extended version, vector storage using Chroma or FAISS can be used to compare current logs with previous issues or known debugging runbooks.",
            ),
            ("p", "Process flow:"),
            (
                "numbered",
                [
                    "The user uploads a system or application log file.",
                    "The backend parses the log and extracts important events.",
                    "The system classifies errors, warnings, and critical messages.",
                    "Relevant log context is passed to the local LLM through Ollama.",
                    "The LLM generates a summary, root cause explanation, and suggested fixes.",
                    "Optional RAG search retrieves similar known issues or previous fixes.",
                    "The dashboard displays analysis, trends, and report output.",
                ],
            ),
        ],
    ),
    (
        "5. Objectives",
        [
            (
                "bullets",
                [
                    "To develop a web-based platform for uploading and analyzing log files.",
                    "To implement log parsing for extracting timestamps, severity levels, error messages, and repeated patterns.",
                    "To integrate Ollama for running a local large language model.",
                    "To generate AI-assisted summaries of important log events.",
                    "To identify possible root causes for errors and system failures.",
                    "To suggest practical debugging actions based on the detected issue.",
                    "To provide a dashboard for viewing error frequency, warnings, and crash trends.",
                    "To explore retrieval-augmented generation for matching logs with known issues and runbooks.",
                    "To maintain privacy by processing logs locally without depending on external cloud APIs.",
                ],
            )
        ],
    ),
]


plan_rows = [
    ("Literature Review & Project Outline", "Week 1 - Week 2", "Study log analysis methods, local LLMs, Ollama, Flask, React, RAG concepts, and vector databases; prepare and submit the project outline."),
    ("System Design & Architecture", "Week 3 - Week 4", "Design the system architecture, upload workflow, parser module, LLM interaction layer, dashboard layout, and database or vector storage structure."),
    ("Frontend & Upload Module", "Week 5 - Week 6", "Build the React, HTML, CSS, and JavaScript interface for uploading logs, viewing analysis results, and navigating dashboard screens."),
    ("Backend & Log Parser Module", "Week 7 - Week 8", "Implement Flask APIs for log upload, parsing, severity classification, pattern extraction, and structured log preparation."),
    ("Local LLM Integration", "Week 9 - Week 10", "Integrate Ollama with suitable local models; create prompts for error summarization, root cause analysis, and suggested debugging actions."),
    ("RAG & Knowledge Base Module", "Week 11 - Week 12", "Implement optional vector search using Chroma or FAISS to retrieve similar past issues, known fixes, and debugging runbooks."),
    ("Dashboard & Report Generation", "Week 13 - Week 14", "Create dashboards for error frequency, warning counts, crash trends, AI summaries, and downloadable log analysis reports."),
    ("Testing, Bug Fixing & Submission", "Week 15 - Week 16", "Perform end-to-end testing with sample logs, fix edge cases, improve UI and backend reliability, finalize documentation, and submit for review."),
]


references = [
    "Xu, W., Huang, L., Fox, A., Patterson, D., & Jordan, M. I. (2009). Detecting Large-Scale System Problems by Mining Console Logs. SOSP. -- Foundational reference for mining system logs and detecting operational problems from log data.",
    "He, S., Zhu, J., He, P., & Lyu, M. R. (2016). Experience Report: System Log Analysis for Anomaly Detection. ISSRE. -- Relevant for understanding how logs can be processed for anomaly detection and failure diagnosis.",
    "Lewis, P., Perez, E., Piktus, A., Petroni, F., Karpukhin, V., Goyal, N., et al. (2020). Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks. NeurIPS. -- Provides the background for using retrieval-based context with language models.",
    "Ollama. (n.d.). Ollama Documentation. https://ollama.com/ -- Official reference for running local large language models used in this project.",
    "Flask. (n.d.). Flask Documentation. https://flask.palletsprojects.com/ -- Official reference for building backend APIs and handling file uploads in Python.",
    "React. (n.d.). React Documentation. https://react.dev/ -- Official reference for developing the frontend interface and dashboard components.",
    "Chroma. (n.d.). Chroma Documentation. https://docs.trychroma.com/ -- Reference for vector storage and retrieval used in the optional RAG module.",
    "FAISS. (n.d.). FAISS Documentation. https://faiss.ai/ -- Reference for efficient similarity search and vector retrieval in log knowledge bases.",
]


supervisor_table = [
    ["", "Supervisor", "Additional Examiner"],
    ["Name", "Thamarai Selvan Gopal", "Thota Rathna Prabha"],
    ["Qualification", "MCA and MBA", "B.Tech"],
    ["Designation", "Senior Technical Lead", "Senior Software Engineer"],
    [
        "Organisation and location",
        "HCLTech, Perungudi OMR Chennai, Tamil Nadu",
        "HCLTech, Ponni Amman Kovil, Sholinganallur OMR Chennai",
    ],
    ["Phone number", "9791143823", "7569521129"],
    ["Email Address", "thamaraiselvan.gopa@hcltech.com", "thota.rathnaprabha@hcltech.com"],
]


remarks = [
    "This project addresses a practical and relevant problem in software testing, system validation, and debugging. Large log files are difficult to review manually, and important failure signals can be missed when errors are spread across multiple components or repeated events.",
    "The proposed solution is technically meaningful because it combines log parsing, local LLM integration, backend development, frontend dashboards, and optional retrieval-augmented generation. Running the LLM through Ollama also makes the project stronger from a privacy and deployment perspective because logs can be analyzed locally.",
    "The project has good academic and engineering value because it is not limited to a simple chatbot. It focuses on extracting structured information from logs, identifying possible causes of failures, and generating useful debugging recommendations for users.",
    "Overall, the project is well scoped and suitable for implementation within the project duration. The student should focus on building a reliable prototype that demonstrates log upload, parsing, AI-based summarization, root cause analysis, suggested fixes, and dashboard-based reporting.",
]


def docx_style(doc):
    s = doc.sections[0]
    s.top_margin = Inches(0.75)
    s.bottom_margin = Inches(0.75)
    s.left_margin = Inches(0.75)
    s.right_margin = Inches(0.75)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(7)
    for style_name, size in [("Title", 18), ("Subtitle", 13), ("Heading 1", 16)]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.font.bold = style_name != "Subtitle"


def border(cell, color="C9C9C9", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell_text(cell, text, bold=False, size=11, white=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    if white:
        r.font.color.rgb = RGBColor(255, 255, 255)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(text)


def build_docx():
    doc = Document()
    docx_style(doc)
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO), width=Inches(0.85))
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(TITLE)
    p = doc.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(SUBTITLE)
    table = doc.add_table(rows=len(metadata), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row, (label, value) in zip(table.rows, metadata):
        cell_text(row.cells[0], label, bold=True)
        cell_text(row.cells[1], value)
        border(row.cells[0], "FFFFFF", "0")
        border(row.cells[1], "FFFFFF", "0")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BIRLA INSTITUTE OF TECHNOLOGY & SCIENCE, PILANI")
    r.bold = True
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("BIRLA INSTITUTE OF TECHNOLOGY & SCIENCE, PILANI").bold = True
    doc.add_heading("Contents", level=1)
    if CONTENTS_IMAGE.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(CONTENTS_IMAGE), width=Inches(5.8))
    doc.add_page_break()
    for heading, blocks in sections:
        doc.add_heading(heading, level=1)
        for kind, val in blocks:
            if kind == "p":
                add_para(doc, val)
            elif kind == "bullets":
                for item in val:
                    doc.add_paragraph(item, style="List Bullet")
            elif kind == "numbered":
                for item in val:
                    doc.add_paragraph(item, style="List Number")
    doc.add_heading("Plan of Work", level=1)
    plan = doc.add_table(rows=1, cols=3)
    for i, h in enumerate(["Phase", "Start - End Date", "Work to be Done"]):
        cell_text(plan.rows[0].cells[i], h, bold=True, white=True)
        shade(plan.rows[0].cells[i], "4F4F4F")
        border(plan.rows[0].cells[i])
    for row_data in plan_rows:
        row = plan.add_row()
        for i, text in enumerate(row_data):
            cell_text(row.cells[i], text, size=10)
            border(row.cells[i])
    doc.add_heading("6. Literature References", level=1)
    for ref in references:
        doc.add_paragraph(ref, style="List Bullet")
    doc.add_heading("7. Particulars of the Supervisor and Examiner", level=1)
    sup = doc.add_table(rows=len(supervisor_table), cols=3)
    for r, row_data in enumerate(supervisor_table):
        for c, text in enumerate(row_data):
            cell_text(sup.rows[r].cells[c], text, bold=(r == 0 or c == 0), size=10)
            if r == 0:
                shade(sup.rows[r].cells[c], "E6E6E6")
            border(sup.rows[r].cells[c])
    doc.add_heading("8. Remarks of the Supervisor", level=1)
    for text in remarks:
        add_para(doc, text)
    sig = doc.add_table(rows=3, cols=3)
    labels = ["Signature of the student", "Signature of the supervisor", "Signature of the Additional Examiner"]
    names = ["Name: Diwakar S", "Name: Thamarai Selvan Gopal", "Name: Thota Rathna Prabha"]
    sign_paths = [STUDENT_SIGN, SUPERVISOR_SIGN, EXAMINER_SIGN]
    for c in range(3):
        cell = sig.rows[0].cells[c]
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if sign_paths[c].exists():
            width = Inches(1.45)
            para.add_run().add_picture(str(sign_paths[c]), width=width)
        else:
            para.add_run("\n\n________________________")
        cell_text(sig.rows[1].cells[c], labels[c], size=11)
        cell_text(sig.rows[2].cells[c], names[c], size=11)
        for r in range(3):
            border(sig.rows[r].cells[c], "FFFFFF", "0")
    doc.save(DOCX_OUT)


def styles():
    base = getSampleStyleSheet()
    base.add(ParagraphStyle("CoverTitle", parent=base["Title"], fontName="Times-Bold", fontSize=18, leading=23, alignment=TA_CENTER, spaceAfter=10))
    base.add(ParagraphStyle("CoverSubtitle", parent=base["Normal"], fontName="Times-Roman", fontSize=13, leading=17, alignment=TA_CENTER, spaceAfter=14))
    base.add(ParagraphStyle("Heading", parent=base["Heading1"], fontName="Times-Bold", fontSize=16, leading=20, spaceBefore=9, spaceAfter=6))
    base.add(ParagraphStyle("Body", parent=base["BodyText"], fontName="Times-Roman", fontSize=11.5, leading=15, alignment=TA_LEFT, spaceAfter=5))
    base.add(ParagraphStyle("BodyCenter", parent=base["Body"], alignment=TA_CENTER))
    base.add(ParagraphStyle("LogBullet", parent=base["Body"], leftIndent=18, firstLineIndent=-10, leading=14, spaceAfter=1))
    base.add(ParagraphStyle("Small", parent=base["Body"], fontSize=9.6, leading=12))
    base.add(ParagraphStyle("SmallBold", parent=base["Small"], fontName="Times-Bold"))
    return base


def esc(text):
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def p(text, style):
    return Paragraph(esc(text), style)


def bullet(text, style):
    return Paragraph("-&nbsp;&nbsp;" + esc(text), style)


def footer(canvas, doc):
    canvas.saveState()
    canvas.setFont("Times-Roman", 8)
    canvas.drawString(0.45 * inch, 0.35 * inch, "Classification: Internal")
    canvas.restoreState()


def build_pdf():
    st = styles()
    doc = BaseDocTemplate(str(PDF_OUT), pagesize=A4, leftMargin=0.58 * inch, rightMargin=0.58 * inch, topMargin=0.62 * inch, bottomMargin=0.62 * inch)
    doc.addPageTemplates([PageTemplate(id="main", frames=[Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height)], onPage=footer)])
    story = []
    story.append(Spacer(1, 0.2 * inch))
    if LOGO.exists():
        img = RLImage(str(LOGO), width=0.82 * inch, height=0.79 * inch)
        img.hAlign = "CENTER"
        story += [img, Spacer(1, 0.16 * inch)]
    story += [Paragraph(TITLE, st["CoverTitle"]), Paragraph(SUBTITLE, st["CoverSubtitle"])]
    meta = [[p(k, st["SmallBold"]), p(v, st["Small"])] for k, v in metadata]
    t = Table(meta, colWidths=[1.9 * inch, 4.55 * inch], hAlign="CENTER")
    t.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "MIDDLE"), ("BOTTOMPADDING", (0, 0), (-1, -1), 4)]))
    story += [t, Spacer(1, 0.4 * inch), Paragraph("<b>BIRLA INSTITUTE OF TECHNOLOGY &amp; SCIENCE, PILANI</b>", st["BodyCenter"]), PageBreak()]
    story += [Paragraph("<b>BIRLA INSTITUTE OF TECHNOLOGY &amp; SCIENCE, PILANI</b>", st["BodyCenter"]), Spacer(1, 0.25 * inch), Paragraph("Contents", st["Heading"])]
    if CONTENTS_IMAGE.exists():
        toc = RLImage(str(CONTENTS_IMAGE), width=6.6 * inch, height=3.55 * inch)
        toc.hAlign = "CENTER"
        story += [Spacer(1, 0.15 * inch), toc]
    story.append(PageBreak())
    for heading, blocks in sections:
        story.append(Paragraph(heading, st["Heading"]))
        for kind, val in blocks:
            if kind == "p":
                story.append(p(val, st["Body"]))
            elif kind == "bullets":
                story.extend(bullet(item, st["LogBullet"]) for item in val)
            elif kind == "numbered":
                story.extend(Paragraph(f"{i}. {esc(item)}", st["LogBullet"]) for i, item in enumerate(val, 1))
    story.append(Paragraph("Plan of Work", st["Heading"]))
    plan_data = [[p("Phase", st["SmallBold"]), p("Start - End Date", st["SmallBold"]), p("Work to be Done", st["SmallBold"])]]
    plan_data += [[p(a, st["Small"]), p(b, st["Small"]), p(c, st["Small"])] for a, b, c in plan_rows]
    plan = Table(plan_data, colWidths=[1.65 * inch, 1.35 * inch, 4.0 * inch], repeatRows=1)
    plan.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4F4F4F")), ("TEXTCOLOR", (0, 0), (-1, 0), colors.white), ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#BFBFBF")), ("VALIGN", (0, 0), (-1, -1), "TOP"), ("TOPPADDING", (0, 0), (-1, -1), 4), ("BOTTOMPADDING", (0, 0), (-1, -1), 4)]))
    story.append(plan)
    story.append(Paragraph("6. Literature References", st["Heading"]))
    story.extend(bullet(ref, st["LogBullet"]) for ref in references)
    story.append(Paragraph("7. Particulars of the Supervisor and Examiner", st["Heading"]))
    sup_data = [[p(cell, st["SmallBold"] if r == 0 or c == 0 else st["Small"]) for c, cell in enumerate(row)] for r, row in enumerate(supervisor_table)]
    sup = Table(sup_data, colWidths=[1.65 * inch, 2.6 * inch, 2.7 * inch], repeatRows=1)
    sup.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E6E6E6")), ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#BFBFBF")), ("VALIGN", (0, 0), (-1, -1), "TOP"), ("TOPPADDING", (0, 0), (-1, -1), 4), ("BOTTOMPADDING", (0, 0), (-1, -1), 4)]))
    story.append(sup)
    story.append(Paragraph("8. Remarks of the Supervisor", st["Heading"]))
    story.extend(p(text, st["Body"]) for text in remarks)
    story.append(Spacer(1, 0.3 * inch))
    def sig_cell(path, label, width=1.45, height=0.45):
        if path.exists():
            return [RLImage(str(path), width=width * inch, height=height * inch), p(label, st["Small"])]
        return [p("________________________", st["Small"]), p(label, st["Small"])]

    sig = Table(
        [
            [
                sig_cell(STUDENT_SIGN, "Signature of the student", 1.45, 0.48),
                sig_cell(SUPERVISOR_SIGN, "Signature of the supervisor", 1.25, 0.50),
                sig_cell(EXAMINER_SIGN, "Signature of the Additional Examiner", 1.45, 0.48),
            ],
            [p("Name: Diwakar S", st["Small"]), p("Name: Thamarai Selvan Gopal", st["Small"]), p("Name: Thota Rathna Prabha", st["Small"])],
        ],
        colWidths=[2.1 * inch, 2.3 * inch, 2.5 * inch],
    )
    sig.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "TOP"), ("TOPPADDING", (0, 0), (-1, -1), 6)]))
    story.append(sig)
    doc.build(story)


if __name__ == "__main__":
    build_docx()
    build_pdf()
