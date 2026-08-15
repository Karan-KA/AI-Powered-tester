import os
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = Path(os.environ.get("REPORT_OUT", ROOT / "generated_reports" / "intelligent-automation-report"))
OUT.mkdir(parents=True, exist_ok=True)

TITLE = "AI-Powered Intelligent Web Test Automation and Analytics Platform"
COURSE = "BITS ZG628T: Dissertation"
MONTH_YEAR = "June 2026"
STUDENT_NAME = "Karan K A"
STUDENT_ID = "202217b2235"
PROGRAMME = "M.Tech. Software Engineering"
ORGANIZATION = "Birla Institute of Technology & Science, Pilani"
SUPERVISOR_NAME = "Gowdeswaran P C"
EXAMINER_NAME = "Aakash R"
ASSET_DIR = Path(os.environ.get("REPORT_ASSET_DIR", r"C:\Users\Acer\AppData\Local\Temp\codex-karan-report-assets"))
BITS_LOGO = ASSET_DIR / "image1.png"
STUDENT_SIGN = ASSET_DIR / "image3.jpg"
SUPERVISOR_SIGN = ASSET_DIR / "image4.jpg"
EXAMINER_SIGN = ASSET_DIR / "image5.jpg"

BLUE = RGBColor(31, 77, 120)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
BORDER = "A6A6A6"
TEXT = RGBColor(25, 25, 25)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color=BORDER, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    values = {"top": top, "start": start, "bottom": bottom, "end": end}
    for key, value in values.items():
        node = margins.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_font(run, size=11, bold=False, italic=False, color=TEXT, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def paragraph(doc, text="", style=None, align=None, before=0, after=6, line=1.10, bold=False, italic=False):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_font(run, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    run = p.add_run(text)
    if level == 1:
        set_font(run, size=16, bold=True, color=BLUE)
    elif level == 2:
        set_font(run, size=13, bold=True, color=BLUE)
    else:
        set_font(run, size=12, bold=True, color=BLUE)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        set_font(run)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        set_font(run)


def add_table(doc, headers, rows, widths=None, caption=None):
    if caption:
        p = paragraph(doc, caption, before=4, after=4, italic=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        set_cell_shading(hdr[i], LIGHT_GRAY)
        set_cell_border(hdr[i])
        set_cell_margins(hdr[i])
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = hdr[i].paragraphs[0].add_run(text)
        set_font(run, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_border(cells[i])
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            run = cells[i].paragraphs[0].add_run(str(text))
            set_font(run, size=10.2)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    return table


def add_page_number(paragraph_obj):
    paragraph_obj.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph_obj.add_run("Page ")
    set_font(run, size=9, color=RGBColor(90, 90, 90))
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def setup_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 12)]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLUE
        style.paragraph_format.space_before = Pt(12 if name == "Heading 2" else 16)
        style.paragraph_format.space_after = Pt(6)


def draw_box(draw, xy, text, fill, outline="#365F91", font=None):
    draw.rounded_rectangle(xy, radius=10, fill=fill, outline=outline, width=2)
    x0, y0, x1, y1 = xy
    lines = text.split("\n")
    total_h = len(lines) * 22
    y = y0 + ((y1 - y0) - total_h) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        draw.text((x0 + ((x1 - x0) - (bbox[2] - bbox[0])) / 2, y), line, fill="#111111", font=font)
        y += 22


def make_architecture_figure():
    path = OUT / "figure-architecture.png"
    img = Image.new("RGB", (1400, 760), "white")
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("arial.ttf", 24)
        small = ImageFont.truetype("arial.ttf", 21)
    except OSError:
        font = ImageFont.load_default()
        small = font
    boxes = [
        ((80, 80, 360, 180), "User Dashboard\nHTML/CSS/JS", "#DDEBF7"),
        ((560, 80, 840, 180), "FastAPI Backend\nREST endpoints", "#E2F0D9"),
        ((1040, 80, 1320, 180), "SQLite Database\nSessions, cases, runs", "#FFF2CC"),
        ((280, 330, 560, 450), "AI Test Generator\nQwen or fallback", "#EADCF8"),
        ((840, 330, 1120, 450), "Playwright Executor\nChromium automation", "#FCE4D6"),
        ((560, 560, 840, 680), "Evidence Store\nScreenshots and logs", "#E7E6E6"),
    ]
    for xy, label, fill in boxes:
        draw_box(draw, xy, label, fill, font=font)
    arrows = [
        ((360, 130), (560, 130)),
        ((840, 130), (1040, 130)),
        ((700, 180), (420, 330)),
        ((700, 180), (980, 330)),
        ((980, 450), (700, 560)),
        ((700, 560), (700, 180)),
    ]
    for start, end in arrows:
        draw.line([start, end], fill="#404040", width=4)
        ex, ey = end
        draw.polygon([(ex, ey), (ex - 14, ey - 8), (ex - 14, ey + 8)], fill="#404040")
    img.save(path)
    return path


def make_workflow_figure():
    path = OUT / "figure-workflow.png"
    img = Image.new("RGB", (1400, 520), "white")
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("arial.ttf", 22)
        small = ImageFont.truetype("arial.ttf", 20)
    except OSError:
        font = ImageFont.load_default()
        small = font
    labels = [
        "Enter URL\nand requirement",
        "Inspect page\ncontrols",
        "Generate\nstructured steps",
        "Validate plan\nand selectors",
        "Execute in\nChromium",
        "Store result\nand evidence",
    ]
    x = 55
    for idx, label in enumerate(labels):
        xy = (x, 140, x + 190, 260)
        draw_box(draw, xy, label, "#DDEBF7" if idx % 2 == 0 else "#E2F0D9", font=font)
        if idx < len(labels) - 1:
            draw.line([(x + 190, 200), (x + 245, 200)], fill="#404040", width=4)
            draw.polygon([(x + 245, 200), (x + 230, 190), (x + 230, 210)], fill="#404040")
        x += 225
    img.save(path)
    return path


def cover_page(doc):
    paragraph(doc, "", after=8)
    if BITS_LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(BITS_LOGO), width=Inches(0.9))
    paragraph(doc, "", after=4)
    p = paragraph(doc, TITLE.upper(), align=WD_ALIGN_PARAGRAPH.CENTER, after=14)
    set_font(p.runs[0], size=18, bold=True, color=RGBColor(0, 0, 0))
    p = paragraph(doc, COURSE, align=WD_ALIGN_PARAGRAPH.CENTER, after=20, bold=True)
    set_font(p.runs[0], size=13, bold=True)
    paragraph(doc, "by", align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    paragraph(doc, STUDENT_NAME, align=WD_ALIGN_PARAGRAPH.CENTER, after=4, bold=True)
    paragraph(doc, STUDENT_ID, align=WD_ALIGN_PARAGRAPH.CENTER, after=24)
    paragraph(doc, "Dissertation work carried out at", align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    paragraph(doc, ORGANIZATION, align=WD_ALIGN_PARAGRAPH.CENTER, after=20)
    paragraph(doc, "Submitted in partial fulfilment of the requirements of", align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    paragraph(doc, f"{PROGRAMME} degree programme", align=WD_ALIGN_PARAGRAPH.CENTER, after=20)
    paragraph(doc, "Under the Supervision of", align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    paragraph(doc, SUPERVISOR_NAME, align=WD_ALIGN_PARAGRAPH.CENTER, after=4, bold=True)
    paragraph(doc, ORGANIZATION, align=WD_ALIGN_PARAGRAPH.CENTER, after=30)
    paragraph(doc, "BIRLA INSTITUTE OF TECHNOLOGY & SCIENCE", align=WD_ALIGN_PARAGRAPH.CENTER, after=4, bold=True)
    paragraph(doc, "PILANI (RAJASTHAN)", align=WD_ALIGN_PARAGRAPH.CENTER, after=14, bold=True)
    paragraph(doc, MONTH_YEAR, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    doc.add_page_break()


def abstract_page(doc):
    add_heading(doc, "ABSTRACT", 1)
    paragraphs = [
        "Frequent releases in web applications make manual regression testing slow, repetitive, and difficult to maintain. This project presents an intelligent automation platform that transforms a user's natural-language testing requirement into executable browser steps. The system combines a FastAPI backend, a responsive dashboard, optional Qwen-based test planning through Ollama, deterministic fallback generation, Playwright browser execution, SQLite persistence, and analytics for reviewing test history.",
        "The platform allows the user to create a saved workspace for a website URL, generate one or more test cases for that workspace, run selected cases in Chromium, and preserve screenshots, logs, status, duration, and error summaries. The design also handles practical constraints found in academic and local environments: when the language model or browser runtime is unavailable, the application still produces useful structured tests or reports warnings instead of failing silently.",
        "The completed implementation demonstrates how intelligent automation can support software quality assurance by reducing script-writing effort while keeping evidence, repeatability, and result visibility. The current backend verification suite contains 34 passing automated tests, and manual testing confirms the main dashboard flow for generating, running, and reviewing web tests.",
    ]
    for text in paragraphs:
        paragraph(doc, text, after=8)
    paragraph(doc, "Keywords: intelligent automation, web testing, Playwright, FastAPI, natural-language test generation, SQLite, quality assurance.", after=20, italic=True)
    sig_table = doc.add_table(rows=3, cols=3)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    labels = ["Signature of the Student", "Signature of the Supervisor", "Signature of the Additional Examiner"]
    images = [STUDENT_SIGN, SUPERVISOR_SIGN, EXAMINER_SIGN]
    names = [f"Name: {STUDENT_NAME}\nID: {STUDENT_ID}\nDate:\nPlace:", f"Name: {SUPERVISOR_NAME}\nDate:\nPlace:", f"Name: {EXAMINER_NAME}\nDate:\nPlace:"]
    for c, value in enumerate(labels):
        cell = sig_table.cell(0, c)
        cell.text = ""
        run = cell.paragraphs[0].add_run(value)
        set_font(run, size=9.5, bold=True)
        set_cell_margins(cell, top=80, bottom=80)
    for c, path in enumerate(images):
        cell = sig_table.cell(1, c)
        cell.text = ""
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if path.exists():
            cell.paragraphs[0].add_run().add_picture(str(path), width=Inches(1.25))
        set_cell_margins(cell, top=80, bottom=80)
    for c, value in enumerate(names):
        cell = sig_table.cell(2, c)
        cell.text = ""
        run = cell.paragraphs[0].add_run(value)
        set_font(run, size=9.5)
        set_cell_margins(cell, top=80, bottom=80)
    for row in sig_table.rows:
        for cell in row.cells:
            set_cell_border(cell, color="FFFFFF", size="0")
    doc.add_page_break()


def contents_page(doc):
    add_heading(doc, "Contents", 1)
    rows = [
        ("1.", "Modules in the Intelligent Automation Platform", "5"),
        ("2.", "Functional Block Diagram and System Description", "8"),
        ("3.", "Major Technical Specifications", "10"),
        ("4.", "Design Considerations", "11"),
        ("5.", "Testing, Results, and Validation", "12"),
        ("6.", "Future Plan", "13"),
        ("7.", "References", "14"),
    ]
    for number, title, page in rows:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        set_font(p.add_run(f"{number} {title} "), bold=False)
        dots = "." * max(8, 70 - len(title))
        set_font(p.add_run(dots + " "), color=RGBColor(90, 90, 90))
        set_font(p.add_run(page), bold=False)
    paragraph(doc, "", after=8)
    add_heading(doc, "List of Figures", 2)
    paragraph(doc, "Figure 1: Logical architecture of the intelligent automation platform .......... 8")
    paragraph(doc, "Figure 2: End-to-end generation and execution workflow ....................... 9")
    add_heading(doc, "List of Tables", 2)
    paragraph(doc, "Table 1: Major software modules .............................................. 5")
    paragraph(doc, "Table 2: Major technical specifications ..................................... 10")
    paragraph(doc, "Table 3: Validation summary ................................................ 12")
    doc.add_page_break()


def modules_section(doc):
    add_heading(doc, "1. MODULES IN THE INTELLIGENT AUTOMATION PLATFORM", 1)
    paragraph(doc, "The project is organized as a set of cooperating modules. Each module has a clear responsibility, which keeps the system understandable for demonstration while still representing a practical automation workflow.")
    add_table(
        doc,
        ["Module", "Implementation", "Responsibility"],
        [
            ("Dashboard", "frontend/dist/index.html", "Provides the browser interface for creating URL workspaces, generating tests, running cases, and reviewing evidence."),
            ("API Layer", "FastAPI routes", "Exposes endpoints for sessions, test generation, execution, saved runs, health status, and analytics."),
            ("AI Test Generator", "ai_service.py and llm_service.py", "Interprets natural-language requirements and creates structured steps using Qwen when available or fallback rules when offline."),
            ("Execution Engine", "executor_service.py", "Runs validated steps through Playwright, handles browser actions, captures screenshots, and classifies outcomes."),
            ("Persistence Layer", "SQLAlchemy and SQLite", "Stores sessions, test cases, run records, logs, screenshots, timestamps, and result metadata."),
            ("Validation Suite", "pytest tests", "Checks API behavior, AI planning rules, selector handling, execution fallbacks, and analytics response shape."),
        ],
        widths=[1.45, 1.9, 3.15],
        caption="Table 1: Major software modules",
    )
    add_heading(doc, "1.1 Dashboard Module", 2)
    paragraph(doc, "The dashboard acts as the user's main workspace. A tester enters a target URL and a requirement such as checking search, login rejection, page load, or exploratory behavior. Generated tests remain attached to the selected URL workspace so earlier prompts are not lost when new tests are added.")
    add_heading(doc, "1.2 Backend API Module", 2)
    paragraph(doc, "The FastAPI backend coordinates all project operations. It serves the static dashboard, initializes the database, validates request and response models through Pydantic schemas, and exposes routes for health checks, sessions, test cases, runs, and analytics.")
    add_heading(doc, "1.3 AI and Fallback Generation Module", 2)
    paragraph(doc, "The generator first inspects the target page to identify visible controls such as inputs, buttons, links, headings, and media elements. If the local language model is available, the requirement is converted into a strict JSON plan. If the model is unavailable or the plan fails validation, the fallback planner creates deterministic steps for common flows such as page loading, search, invalid sign-in, and evidence capture.")
    add_heading(doc, "1.4 Playwright Execution Module", 2)
    paragraph(doc, "The execution service opens Chromium, performs actions such as navigation, click, fill, key press, wait, and assertion checks, and records screenshots and logs. Optional steps are treated carefully so that missing non-essential UI elements produce warnings rather than unnecessary hard failures.")
    add_heading(doc, "1.5 Database and Analytics Module", 2)
    paragraph(doc, "SQLite is used for a lightweight academic deployment, while SQLAlchemy keeps the model portable for later migration. Analytics are calculated from saved run records, including total cases, total runs, pass rate, warning count, failure count, average duration, and recent executions.")
    doc.add_page_break()


def functional_section(doc, arch_path, workflow_path):
    add_heading(doc, "2. FUNCTIONAL BLOCK DIAGRAM AND SYSTEM DESCRIPTION", 1)
    paragraph(doc, "The platform follows a request-to-evidence workflow. The user starts from a web dashboard, the backend prepares a structured test plan, and Playwright executes that plan in a real browser. The result is stored with enough context for later review.")
    doc.add_picture(str(arch_path), width=Inches(6.3))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph(doc, "Figure 1: Logical architecture of the intelligent automation platform", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
    paragraph(doc, "The dashboard sends a generation request to FastAPI. The backend records or reuses a URL workspace, asks the generator for test steps, stores the generated case, and returns the case to the frontend. During execution, Playwright interacts with the selected website and the backend stores screenshots, console logs, step logs, duration, status, and error summaries.")
    doc.add_picture(str(workflow_path), width=Inches(6.3))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph(doc, "Figure 2: End-to-end generation and execution workflow", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
    add_heading(doc, "2.1 Data Flow", 2)
    add_numbered(
        doc,
        [
            "The user creates or opens a URL workspace from the dashboard.",
            "The user enters a requirement, or leaves the prompt blank for an exploratory suite.",
            "The backend inspects the page and identifies visible controls.",
            "The AI planner creates structured steps; fallback rules are used when the model is unavailable or invalid.",
            "The selected test case is executed in Chromium using Playwright.",
            "Logs, screenshots, status, duration, and error details are saved in SQLite.",
            "The dashboard displays generated steps, execution evidence, and analytics.",
        ],
    )
    add_heading(doc, "2.2 Status Classification", 2)
    paragraph(doc, "A run is marked passed when all required steps complete, failed when a required step cannot be completed, warning when optional steps or environment dependencies are missing, and running while execution is still active.")
    doc.add_page_break()


def specs_section(doc):
    add_heading(doc, "3. MAJOR TECHNICAL SPECIFICATIONS", 1)
    add_table(
        doc,
        ["Category", "Specification"],
        [
            ("Project type", "AI-assisted intelligent web test automation and analytics platform"),
            ("Frontend", "Responsive HTML, CSS, and JavaScript dashboard"),
            ("Backend framework", "FastAPI with Pydantic validation"),
            ("Automation runtime", "Playwright with Chromium browser execution"),
            ("AI component", "Ollama with Qwen for local planning, plus deterministic fallback generation"),
            ("Database", "SQLite through SQLAlchemy ORM"),
            ("Stored entities", "URL workspaces, generated test cases, test runs, logs, screenshots, status, timing, and analytics data"),
            ("Supported actions", "goto, click, fill, press, wait, screenshot, title assertion, text assertion, visibility assertion, URL assertion, value assertion, and count assertion"),
            ("Verification", "34 automated backend tests passing through pytest"),
            ("Deployment mode", "Local academic demonstration with scope for later server deployment"),
        ],
        widths=[2.1, 4.4],
        caption="Table 2: Major technical specifications",
    )
    add_heading(doc, "3.1 Software Environment", 2)
    add_bullets(
        doc,
        [
            "Python 3.11 is used for the backend application and automated tests.",
            "FastAPI provides REST endpoints and interactive Swagger documentation.",
            "Playwright drives browser automation and screenshot capture.",
            "SQLite keeps the setup simple for local execution and viva demonstration.",
            "Ollama/Qwen integration is optional, allowing the system to work even when the model is not running.",
        ],
    )
    add_heading(doc, "3.2 API Surface", 2)
    paragraph(doc, "The main endpoints include health checks, session creation and listing, test generation, suite generation, test retrieval, run execution, run history, clearing saved records, and analytics. This separation supports both frontend use and independent API testing.")
    doc.add_page_break()


def design_section(doc):
    add_heading(doc, "4. DESIGN CONSIDERATIONS", 1)
    paragraph(doc, "The design gives importance to reliability, explainability, and local usability. AI-generated automation can be helpful, but it must be constrained because invalid browser actions or invented selectors would reduce trust. For this reason, the project uses structured schemas, page inspection, supported action lists, validation, fallback planning, and runtime selector recovery.")
    add_heading(doc, "4.1 Reliability and Fallback Behavior", 2)
    paragraph(doc, "The system does not depend entirely on the language model. If Qwen is unavailable or returns a plan that does not satisfy the expected schema, deterministic rules create a useful test case. This makes the application demonstrable on machines where the local model is not installed or temporarily offline.")
    add_heading(doc, "4.2 Evidence Preservation", 2)
    paragraph(doc, "Testing is useful only when the result can be reviewed. Each run stores logs, screenshots, status, duration, and error summaries. This gives the user a traceable record instead of only a pass/fail message.")
    add_heading(doc, "4.3 Safety of Automated Actions", 2)
    paragraph(doc, "The AI instructions and fallback flows avoid destructive actions such as deleting data, placing orders, or using real credentials. Authentication tests use deliberately invalid accounts and verify rejection instead of attempting real login.")
    add_heading(doc, "4.4 Academic Demonstration Scope", 2)
    paragraph(doc, "SQLite and local execution were selected to reduce setup effort. The architecture remains modular, so PostgreSQL, authentication, CI/CD integration, and parallel execution can be added later without rewriting the complete system.")
    doc.add_page_break()


def testing_section(doc):
    add_heading(doc, "5. TESTING, RESULTS, AND VALIDATION", 1)
    paragraph(doc, "The project was validated through automated backend tests and manual dashboard checks. The automated tests cover the API, AI service behavior, and execution service fallbacks. Manual checks cover the visible user journey from creating a workspace to reviewing generated tests and run evidence.")
    add_table(
        doc,
        ["Validation Area", "Coverage", "Observed Result"],
        [
            ("API health and schemas", "Health endpoint and request-response models", "Responses are shaped correctly."),
            ("Test generation", "Focused prompts, blank prompts, saved cases, and suite generation", "Generated tests are stored and returned to the dashboard."),
            ("Session behavior", "Workspace creation, listing, update, deletion, and session-scoped histories", "URL workspaces preserve related test and run records."),
            ("AI planning rules", "Intent extraction, plan validation, selector grounding, and fallback use", "Invalid or unavailable AI output is handled safely."),
            ("Execution service", "Playwright availability, selector recovery, screenshot logic, logs, warnings, and failures", "Runs classify outcomes and preserve evidence."),
            ("Analytics", "Counts, pass rate, warnings, failures, average duration, and recent runs", "Dashboard-ready analytics are returned."),
        ],
        widths=[1.65, 2.65, 2.2],
        caption="Table 3: Validation summary",
    )
    add_heading(doc, "5.1 Automated Test Result", 2)
    paragraph(doc, "The current backend verification result is 34 passed tests. This result confirms that the repository's main backend behaviors are internally consistent at the time this report was prepared.")
    add_heading(doc, "5.2 Result Summary", 2)
    paragraph(doc, "The implemented system can create a URL workspace, generate test cases from natural language or fallback logic, execute selected cases in a browser, store evidence, and display analytics. The result satisfies the project objective of demonstrating an intelligent automation workflow that reduces manual script-writing effort while keeping the test process auditable.")
    doc.add_page_break()


def future_section(doc):
    add_heading(doc, "6. FUTURE PLAN", 1)
    add_bullets(
        doc,
        [
            "Add user authentication and role-based access for multi-user operation.",
            "Move from SQLite to PostgreSQL for deployed environments.",
            "Introduce stronger locator healing using historical run evidence and DOM comparison.",
            "Support CI/CD execution so generated tests can run automatically during releases.",
            "Add downloadable PDF or HTML execution reports for each test run.",
            "Implement parallel browser execution for faster suite runs.",
            "Add visual regression checks by comparing screenshots across builds.",
            "Expand analytics with trend charts, flaky test identification, and failure clustering.",
        ],
    )
    add_heading(doc, "7. REFERENCES", 1)
    refs = [
        "FastAPI documentation, https://fastapi.tiangolo.com/",
        "Playwright documentation, https://playwright.dev/python/",
        "SQLAlchemy documentation, https://www.sqlalchemy.org/",
        "SQLite documentation, https://www.sqlite.org/docs.html",
        "Ollama documentation, https://ollama.com/",
        "Project repository documentation: README.md, ARCHITECTURE.md, API_REFERENCE.md, TESTING.md, and USER_MANUAL.md.",
    ]
    add_numbered(doc, refs)


def main():
    arch = make_architecture_figure()
    workflow = make_workflow_figure()
    doc = Document()
    setup_styles(doc)
    footer = doc.sections[0].footer.paragraphs[0]
    add_page_number(footer)
    cover_page(doc)
    abstract_page(doc)
    contents_page(doc)
    modules_section(doc)
    functional_section(doc, arch, workflow)
    specs_section(doc)
    design_section(doc)
    testing_section(doc)
    future_section(doc)

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.49)
        section.footer_distance = Inches(0.49)

    output = OUT / "Intelligent_Automation_Mid_Sem_Report.docx"
    doc.save(output)
    print(output)


if __name__ == "__main__":
    main()
