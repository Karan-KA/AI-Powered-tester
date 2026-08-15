import os
from pathlib import Path

from PIL import Image as PILImage, ImageDraw, ImageFont, ImageOps
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


OUT = Path(os.environ.get("REPORT_OUT", r"C:\Users\Acer\AppData\Local\Temp\codex-diwakar-midsem-report"))
ASSETS = Path(os.environ.get("REPORT_ASSET_DIR", r"C:\Users\Acer\AppData\Local\Temp\codex-diwakar-report-assets"))
OUT.mkdir(parents=True, exist_ok=True)

TITLE = "Secure Offline Log Analysis using Local LLM (AI)"
SUBTITLE = "An AI-Driven System for Local Log Processing, Root Cause Analysis, and Debugging Recommendations"
COURSE = "BITS ZG628T: Dissertation"
STUDENT_NAME = "Diwakar S"
STUDENT_ID = "202217b2100"
PROGRAMME = "BSc - Design and Computing-HCL"
ORG = "HCLTech"
SUPERVISOR = "Thamarai Selvan Gopal"
EXAMINER = "Thota Rathna Prabha"
MONTH_YEAR = "June 2026"

BITS_LOGO = ASSETS / "image1.png"
STUDENT_SIGN = ASSETS / "image3.jpg"
SUPERVISOR_SIGN_RAW = ASSETS / "image4.jpg"
EXAMINER_SIGN_RAW = ASSETS / "image5.jpg"
SUPERVISOR_SIGN = OUT / "supervisor-sign-cropped.jpg"
EXAMINER_SIGN = OUT / "examiner-sign-cropped.jpg"

BLUE = RGBColor(31, 77, 120)
TEXT = RGBColor(20, 20, 20)


def crop_signature(src: Path, dst: Path):
    if not src.exists():
        return src
    img = PILImage.open(src).convert("RGB")
    gray = ImageOps.grayscale(img)
    mask = gray.point(lambda p: 255 if p < 210 else 0)
    bbox = mask.getbbox()
    if bbox:
        left, top, right, bottom = bbox
        pad = 25
        img = img.crop((max(0, left - pad), max(0, top - pad), min(img.width, right + pad), min(img.height, bottom + pad)))
    img.save(dst, quality=92)
    return dst


def set_font(run, size=11, bold=False, italic=False, color=TEXT, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def para(doc, text="", align=None, before=0, after=6, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_font(r, size=size, bold=bold, italic=italic)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    r = p.add_run(text)
    set_font(r, size=16 if level == 1 else 13, bold=True, color=BLUE)


def cell_border(cell, color="A6A6A6", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:color"), color)


def cell_shade(cell, fill):
    pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    pr.append(shd)


def add_docx_table(doc, headers, rows, widths=None, caption=None):
    if caption:
        para(doc, caption, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, after=4)
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell_shade(cell, "F2F4F7")
        cell_border(cell)
        run = cell.paragraphs[0].add_run(h)
        set_font(run, bold=True, size=10)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cell_border(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            run = cells[i].paragraphs[0].add_run(str(value))
            set_font(run, size=9.5)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    return t


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        set_font(p.add_run(item), size=10.7)


def make_figures():
    try:
        font = ImageFont.truetype("arial.ttf", 24)
    except OSError:
        font = ImageFont.load_default()

    def box(draw, xy, text, fill):
        draw.rounded_rectangle(xy, radius=10, fill=fill, outline="#365F91", width=2)
        x0, y0, x1, y1 = xy
        lines = text.split("\n")
        y = y0 + ((y1 - y0) - len(lines) * 24) / 2
        for line in lines:
            b = draw.textbbox((0, 0), line, font=font)
            draw.text((x0 + ((x1 - x0) - (b[2] - b[0])) / 2, y), line, fill="#111", font=font)
            y += 24

    arch = OUT / "diwakar-architecture.png"
    img = PILImage.new("RGB", (1400, 760), "white")
    d = ImageDraw.Draw(img)
    boxes = [
        ((70, 90, 350, 190), "React Dashboard\nUpload and results", "#DDEBF7"),
        ((550, 90, 850, 190), "Flask Backend\nAPIs and workflow", "#E2F0D9"),
        ((1050, 90, 1330, 190), "Local Storage\nReports and history", "#FFF2CC"),
        ((230, 330, 510, 450), "Log Parser\nPatterns and severity", "#EADCF8"),
        ((590, 330, 870, 450), "Ollama Local LLM\nSummary and RCA", "#FCE4D6"),
        ((950, 330, 1230, 450), "RAG Knowledge Base\nKnown issue match", "#E7E6E6"),
        ((590, 560, 870, 680), "Dashboard Output\nTrends and fixes", "#D9EAD3"),
    ]
    for xy, text, fill in boxes:
        box(d, xy, text, fill)
    for start, end in [((350, 140), (550, 140)), ((850, 140), (1050, 140)), ((700, 190), (370, 330)), ((700, 190), (730, 330)), ((700, 190), (1090, 330)), ((730, 450), (730, 560))]:
        d.line([start, end], fill="#404040", width=4)
        d.ellipse((end[0] - 5, end[1] - 5, end[0] + 5, end[1] + 5), fill="#404040")
    img.save(arch)

    flow = OUT / "diwakar-workflow.png"
    img = PILImage.new("RGB", (1400, 460), "white")
    d = ImageDraw.Draw(img)
    labels = ["Upload\nlog file", "Parse events\nand severity", "Extract\ncontext", "Analyze with\nlocal LLM", "Match known\nissues", "Show RCA\nand fixes"]
    x = 50
    for i, label in enumerate(labels):
        box(d, (x, 150, x + 190, 260), label, "#DDEBF7" if i % 2 == 0 else "#E2F0D9")
        if i < len(labels) - 1:
            d.line([(x + 190, 205), (x + 235, 205)], fill="#404040", width=4)
            d.polygon([(x + 235, 205), (x + 220, 195), (x + 220, 215)], fill="#404040")
        x += 225
    img.save(flow)
    return arch, flow


def build_docx():
    sup = crop_signature(SUPERVISOR_SIGN_RAW, SUPERVISOR_SIGN)
    ex = crop_signature(EXAMINER_SIGN_RAW, EXAMINER_SIGN)
    arch, flow = make_figures()
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    p = para(doc, "", after=8)
    if BITS_LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(BITS_LOGO), width=Inches(0.85))
    p = para(doc, TITLE.upper(), align=WD_ALIGN_PARAGRAPH.CENTER, after=4, bold=True, size=17)
    para(doc, SUBTITLE, align=WD_ALIGN_PARAGRAPH.CENTER, after=14, size=11)
    para(doc, COURSE, align=WD_ALIGN_PARAGRAPH.CENTER, after=18, bold=True)
    para(doc, "by", align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, STUDENT_NAME, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    para(doc, STUDENT_ID, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
    para(doc, "Dissertation work carried out at", align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, ORG, align=WD_ALIGN_PARAGRAPH.CENTER, after=16)
    para(doc, f"Submitted in partial fulfilment of {PROGRAMME} degree programme", align=WD_ALIGN_PARAGRAPH.CENTER, after=16)
    para(doc, "Under the Supervision of", align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, SUPERVISOR, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, after=26)
    para(doc, "BIRLA INSTITUTE OF TECHNOLOGY & SCIENCE", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    para(doc, "PILANI (RAJASTHAN)", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    para(doc, MONTH_YEAR, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    heading(doc, "ABSTRACT")
    for text in [
        "Modern validation environments and software services generate large log files that contain errors, warnings, timestamps, component messages, and crash signals. Manual review of these files is slow and often depends on the experience of the engineer reading them. This project proposes a secure offline log analysis platform that processes logs locally and uses a local large language model to explain failures in a readable form.",
        "The system accepts uploaded log files, parses important events, classifies severity levels, extracts relevant context, and sends only local context to an Ollama-based model. The output includes an error summary, probable root cause, debugging recommendations, and optional matching with known issues through retrieval-augmented generation. Since processing is performed on the user's machine, sensitive log data does not need to be sent to an external cloud API.",
        "The project demonstrates a practical artificial-intelligence workflow for software testing, validation, and support teams. It combines Python-based parsing, Flask APIs, a React dashboard, local LLM inference, optional vector search, and visual reporting to reduce the time required for log review and initial debugging.",
    ]:
        para(doc, text, after=8)
    para(doc, "Keywords: offline log analysis, local LLM, Ollama, Flask, React, root cause analysis, RAG, debugging recommendations.", italic=True, after=18)
    t = doc.add_table(rows=3, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    labels = ["Signature of the Student", "Signature of the Supervisor", "Signature of the Additional Examiner"]
    imgs = [STUDENT_SIGN, sup, ex]
    names = [f"Name: {STUDENT_NAME}\nID: {STUDENT_ID}\nDate:\nPlace:", f"Name: {SUPERVISOR}\nDate:\nPlace:", f"Name: {EXAMINER}\nDate:\nPlace:"]
    for c in range(3):
        set_font(t.cell(0, c).paragraphs[0].add_run(labels[c]), bold=True, size=9)
        t.cell(1, c).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if imgs[c].exists():
            t.cell(1, c).paragraphs[0].add_run().add_picture(str(imgs[c]), width=Inches(1.3))
        set_font(t.cell(2, c).paragraphs[0].add_run(names[c]), size=9)
    doc.add_page_break()

    heading(doc, "Contents")
    for line in [
        "1. Modules in Secure Offline Log Analysis Platform ................................ 5",
        "2. Functional Block Diagram and System Description ............................... 8",
        "3. Major Technical Specifications ................................................ 9",
        "4. Design Considerations ......................................................... 10",
        "5. Future Plan ................................................................... 10",
        "Figure 1: Logical architecture of the offline log analysis platform ............... 8",
        "Figure 2: Log analysis workflow .................................................. 9",
        "Table 1: Major software modules .................................................. 5",
        "Table 2: Technical specifications ................................................ 9",
    ]:
        para(doc, line, after=4)
    doc.add_page_break()

    heading(doc, "1. MODULES IN SECURE OFFLINE LOG ANALYSIS PLATFORM")
    para(doc, "The proposed system is divided into modules that handle log upload, parsing, local analysis, retrieval, dashboard display, and report preparation.")
    add_docx_table(doc, ["Module", "Technology", "Responsibility"], [
        ("Frontend dashboard", "React, HTML, CSS, JavaScript", "Provides log upload, result viewing, charts, and navigation."),
        ("Backend API", "Flask", "Handles upload requests, invokes parser services, and returns structured analysis."),
        ("Log parser", "Python", "Extracts timestamps, severity levels, errors, warnings, repeated patterns, and important context."),
        ("Local LLM layer", "Ollama", "Generates summaries, probable root causes, and debugging recommendations locally."),
        ("RAG module", "Chroma or FAISS", "Retrieves related known issues, runbooks, or previous fixes for context enrichment."),
        ("Reporting module", "Dashboard and exports", "Displays trends, error frequency, RCA, and suggested actions."),
    ], widths=[1.6, 1.8, 3.1], caption="Table 1: Major software modules")
    heading(doc, "1.1 Description of Major Modules", 2)
    bullets(doc, [
        "The upload module accepts logs from software, driver, service, or validation environments.",
        "The parser removes noise and organizes important lines into structured events.",
        "The LLM module prepares human-readable explanations while keeping data local.",
        "The dashboard converts parsed and AI-generated output into summaries, tables, and visual trends.",
    ])
    doc.add_page_break()

    heading(doc, "2. FUNCTIONAL BLOCK DIAGRAM AND SYSTEM DESCRIPTION")
    para(doc, "The functional architecture begins with a user uploading a log file. The Flask backend stores the request, parses the file, prepares important context, sends it to the local LLM, optionally searches known issues, and returns a structured analysis to the dashboard.")
    doc.add_picture(str(arch), width=Inches(6.3))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc, "Figure 1: Logical architecture of the offline log analysis platform", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
    doc.add_picture(str(flow), width=Inches(6.3))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc, "Figure 2: Log analysis workflow", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
    doc.add_page_break()

    heading(doc, "3. MAJOR TECHNICAL SPECIFICATIONS")
    add_docx_table(doc, ["Category", "Specification"], [
        ("Project type", "Secure offline AI-assisted log analysis and debugging platform"),
        ("Backend", "Python Flask APIs for upload, parsing, and analysis workflow"),
        ("Frontend", "React with HTML, CSS, and JavaScript dashboard"),
        ("AI runtime", "Ollama with a locally hosted large language model"),
        ("Retrieval option", "Chroma or FAISS vector database for known issue matching"),
        ("Primary outputs", "Log summary, error classification, probable root cause, suggested fixes, and dashboard trends"),
        ("Privacy model", "Local processing without sending logs to cloud AI services"),
    ], widths=[2.1, 4.4], caption="Table 2: Technical specifications")
    heading(doc, "4. DESIGN CONSIDERATIONS")
    para(doc, "The system is designed for privacy, clarity, and practical debugging support. Local processing protects sensitive logs, while structured parsing prevents the LLM from receiving unnecessary noisy text. The dashboard is intended to help engineers move from raw logs to actionable debugging notes quickly.")
    heading(doc, "5. FUTURE PLAN")
    bullets(doc, [
        "Improve parser support for more log formats and timestamp styles.",
        "Add stronger anomaly detection and clustering for repeated failures.",
        "Expand the known-issue knowledge base with runbooks and historical fixes.",
        "Generate downloadable PDF or HTML analysis reports.",
        "Add user authentication and project-level log history.",
    ])
    heading(doc, "6. REFERENCES")
    bullets(doc, [
        "Ollama Documentation, https://ollama.com/",
        "Flask Documentation, https://flask.palletsprojects.com/",
        "React Documentation, https://react.dev/",
        "Chroma Documentation, https://docs.trychroma.com/",
        "FAISS Documentation, https://faiss.ai/",
        "Lewis et al. (2020), Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks.",
        "He et al. (2016), Experience Report: System Log Analysis for Anomaly Detection.",
    ])

    path = OUT / "Diwakar_Secure_Offline_Log_Analysis_Mid_Sem_Report.docx"
    doc.save(path)
    return path, arch, flow, sup, ex


def pdf_styles():
    s = getSampleStyleSheet()
    s["Normal"].fontName = "Helvetica"
    s["Normal"].fontSize = 10.5
    s["Normal"].leading = 14
    s["Heading1"].fontName = "Helvetica-Bold"
    s["Heading1"].fontSize = 15
    s["Heading1"].textColor = colors.HexColor("#1F4D78")
    s["Heading2"].fontName = "Helvetica-Bold"
    s["Heading2"].fontSize = 12
    s["Heading2"].textColor = colors.HexColor("#1F4D78")
    s.add(ParagraphStyle("Center", parent=s["Normal"], alignment=TA_CENTER))
    s.add(ParagraphStyle("Title2", parent=s["Title"], fontName="Helvetica-Bold", fontSize=18, leading=23, alignment=TA_CENTER))
    s.add(ParagraphStyle("Small", parent=s["Normal"], fontSize=8.5, leading=10))
    return s


def P(text, style):
    return Paragraph(text, style)


def pdf_table(data, widths):
    cell = ParagraphStyle("Cell", fontName="Helvetica", fontSize=8.2, leading=10)
    head = ParagraphStyle("Head", fontName="Helvetica-Bold", fontSize=8.5, leading=10)
    wrapped = []
    for r, row in enumerate(data):
        st = head if r == 0 else cell
        wrapped.append([x if hasattr(x, "drawOn") else Paragraph(str(x), st) for x in row])
    t = Table(wrapped, colWidths=widths, repeatRows=1)
    t.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#A6A6A6")), ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")), ("VALIGN", (0, 0), (-1, -1), "MIDDLE"), ("LEFTPADDING", (0, 0), (-1, -1), 6), ("RIGHTPADDING", (0, 0), (-1, -1), 6), ("TOPPADDING", (0, 0), (-1, -1), 5), ("BOTTOMPADDING", (0, 0), (-1, -1), 5)]))
    return t


def build_pdf(arch, flow, sup, ex):
    s = pdf_styles()
    story = []
    if BITS_LOGO.exists():
        story += [Image(str(BITS_LOGO), width=0.8 * inch, height=0.77 * inch), Spacer(1, 0.16 * inch)]
    story += [
        P(TITLE.upper(), s["Title2"]), P(SUBTITLE, s["Center"]), Spacer(1, 0.12 * inch),
        P(COURSE, s["Center"]), Spacer(1, 0.15 * inch), P("by", s["Center"]), P(STUDENT_NAME, s["Center"]), P(STUDENT_ID, s["Center"]),
        Spacer(1, 0.15 * inch), P("Dissertation work carried out at", s["Center"]), P(ORG, s["Center"]),
        Spacer(1, 0.15 * inch), P(f"Submitted in partial fulfilment of {PROGRAMME} degree programme", s["Center"]),
        Spacer(1, 0.15 * inch), P("Under the Supervision of", s["Center"]), P(SUPERVISOR, s["Center"]),
        Spacer(1, 0.35 * inch), P("BIRLA INSTITUTE OF TECHNOLOGY & SCIENCE", s["Center"]), P("PILANI (RAJASTHAN)", s["Center"]), P(MONTH_YEAR, s["Center"]), PageBreak(),
        P("ABSTRACT", s["Heading1"]),
        P("Modern validation environments and software services generate large log files that contain errors, warnings, timestamps, component messages, and crash signals. Manual review of these files is slow and often depends on the experience of the engineer reading them. This project proposes a secure offline log analysis platform that processes logs locally and uses a local large language model to explain failures in a readable form.", s["Normal"]),
        P("The system accepts uploaded log files, parses important events, classifies severity levels, extracts relevant context, and sends only local context to an Ollama-based model. The output includes an error summary, probable root cause, debugging recommendations, and optional matching with known issues through retrieval-augmented generation.", s["Normal"]),
        P("Keywords: offline log analysis, local LLM, Ollama, Flask, React, root cause analysis, RAG, debugging recommendations.", s["Normal"]), Spacer(1, 0.25 * inch),
        pdf_table([["Signature of the Student", "Signature of the Supervisor", "Signature of the Additional Examiner"], [Image(str(STUDENT_SIGN), 1.25 * inch, 0.48 * inch), Image(str(sup), 1.0 * inch, 0.48 * inch), Image(str(ex), 1.25 * inch, 0.48 * inch)], [f"Name: {STUDENT_NAME}<br/>ID: {STUDENT_ID}<br/>Date:<br/>Place:", f"Name: {SUPERVISOR}<br/>Date:<br/>Place:", f"Name: {EXAMINER}<br/>Date:<br/>Place:"]], [2 * inch, 2 * inch, 2 * inch]), PageBreak(),
        P("Contents", s["Heading1"]),
        P("1. Modules in Secure Offline Log Analysis Platform ............................ 4", s["Normal"]),
        P("2. Functional Block Diagram and System Description ........................... 5", s["Normal"]),
        P("3. Major Technical Specifications ............................................ 6", s["Normal"]),
        P("4. Design Considerations ..................................................... 7", s["Normal"]),
        P("5. Future Plan ............................................................... 7", s["Normal"]),
        P("Figure 1: Logical architecture of the offline log analysis platform", s["Normal"]),
        P("Figure 2: Log analysis workflow", s["Normal"]), PageBreak(),
        P("1. MODULES IN SECURE OFFLINE LOG ANALYSIS PLATFORM", s["Heading1"]),
        P("The proposed system is divided into modules that handle log upload, parsing, local analysis, retrieval, dashboard display, and report preparation.", s["Normal"]),
        pdf_table([["Module", "Technology", "Responsibility"], ["Frontend dashboard", "React, HTML, CSS, JavaScript", "Provides log upload, result viewing, charts, and navigation."], ["Backend API", "Flask", "Handles upload requests, parser services, and structured analysis."], ["Log parser", "Python", "Extracts timestamps, severity, errors, warnings, and important context."], ["Local LLM layer", "Ollama", "Generates summaries, probable root causes, and debugging recommendations locally."], ["RAG module", "Chroma or FAISS", "Retrieves related known issues, runbooks, or previous fixes."], ["Reporting module", "Dashboard and exports", "Displays trends, error frequency, RCA, and suggested actions."]], [1.55 * inch, 1.8 * inch, 2.85 * inch]),
        P("Table 1: Major software modules", s["Center"]), PageBreak(),
        P("2. FUNCTIONAL BLOCK DIAGRAM AND SYSTEM DESCRIPTION", s["Heading1"]),
        P("The functional architecture begins with a user uploading a log file. The backend parses the file, prepares context, sends it to the local LLM, optionally searches known issues, and returns structured analysis.", s["Normal"]),
        Image(str(arch), 6.2 * inch, 3.35 * inch), P("Figure 1: Logical architecture of the offline log analysis platform", s["Center"]),
        Image(str(flow), 6.2 * inch, 2.0 * inch), P("Figure 2: Log analysis workflow", s["Center"]), PageBreak(),
        P("3. MAJOR TECHNICAL SPECIFICATIONS", s["Heading1"]),
        pdf_table([["Category", "Specification"], ["Project type", "Secure offline AI-assisted log analysis and debugging platform"], ["Backend", "Python Flask APIs for upload, parsing, and analysis workflow"], ["Frontend", "React with HTML, CSS, and JavaScript dashboard"], ["AI runtime", "Ollama with a locally hosted large language model"], ["Retrieval option", "Chroma or FAISS vector database"], ["Primary outputs", "Log summary, error classification, probable root cause, suggested fixes, and dashboard trends"], ["Privacy model", "Local processing without cloud AI services"]], [2.0 * inch, 4.2 * inch]),
        P("Table 2: Technical specifications", s["Center"]),
        P("4. DESIGN CONSIDERATIONS", s["Heading1"]),
        P("The system is designed for privacy, clarity, and practical debugging support. Local processing protects sensitive logs, while structured parsing prevents the LLM from receiving unnecessary noisy text.", s["Normal"]),
        P("5. FUTURE PLAN", s["Heading1"]),
        P("- Improve parser support for more log formats and timestamp styles.<br/>- Add stronger anomaly detection and clustering.<br/>- Expand the known-issue knowledge base with runbooks and historical fixes.<br/>- Generate downloadable PDF or HTML analysis reports.<br/>- Add user authentication and project-level log history.", s["Normal"]),
        P("6. REFERENCES", s["Heading1"]),
        P("- Ollama Documentation, https://ollama.com/<br/>- Flask Documentation, https://flask.palletsprojects.com/<br/>- React Documentation, https://react.dev/<br/>- Chroma Documentation, https://docs.trychroma.com/<br/>- FAISS Documentation, https://faiss.ai/<br/>- Lewis et al. (2020), Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks.<br/>- He et al. (2016), Experience Report: System Log Analysis for Anomaly Detection.", s["Normal"]),
    ]
    path = OUT / "Diwakar_Secure_Offline_Log_Analysis_Mid_Sem_Report.pdf"
    doc = SimpleDocTemplate(str(path), pagesize=letter, rightMargin=inch, leftMargin=inch, topMargin=inch, bottomMargin=inch)
    doc.build(story)
    return path


if __name__ == "__main__":
    docx, arch, flow, sup, ex = build_docx()
    pdf = build_pdf(arch, flow, sup, ex)
    print(docx)
    print(pdf)
